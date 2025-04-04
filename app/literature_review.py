from fastapi import FastAPI, HTTPException, APIRouter
from pydantic import BaseModel
from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

app = FastAPI()
router = APIRouter()

# Allow CORS for your frontend domain
origins = [
    "http://localhost:5173",  # Replace with your frontend URL
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],  # Allow all HTTP methods
    allow_headers=["*"],  # Allow all headers
)

# Load the model and tokenizer
model_name = "facebook/bart-large-cnn"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForSeq2SeqLM.from_pretrained(model_name)

# Input model for summarizing papers
class PaperRequest(BaseModel):
    topic: str
    num_papers: int
    papers: list  # List of dictionaries containing title, authors, year, and abstract

class SummarizedPaper(BaseModel):
    title: str
    authors: str
    year: int
    summary: str  # Updated to match the field in the frontend

# Function to summarize the paper abstract
def summarize_text(text: str) -> str:
    inputs = tokenizer.encode(text, return_tensors="pt", max_length=1024, truncation=True)
    summary_ids = model.generate(inputs, max_length=150, min_length=30, length_penalty=2.0, num_beams=4, early_stopping=True)
    summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
    return summary

# Route for summarizing papers
@router.post("/api/fetch-summarized-papers")
async def fetch_summarized_papers(request: PaperRequest):
    try:
        if not request.papers or len(request.papers) == 0:
            raise HTTPException(status_code=400, detail="No papers provided for summarization")

        summarized_papers = []
        for paper in request.papers:
            if "abstract" not in paper:
                raise HTTPException(status_code=400, detail="Missing abstract in paper data")
            summarized_abstract = summarize_text(paper["abstract"])
            summarized_papers.append(SummarizedPaper(
                title=paper["title"],
                authors=paper["authors"],
                year=paper["year"],
                summary=summarized_abstract  # This is now 'summary' to match frontend
            ))

        return {"papers": summarized_papers}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Input model for generating a downloadable DOCX file
class DownloadDocRequest(BaseModel):
    topic: str
    papers: list[SummarizedPaper]

# Route to download a DOCX file with summarized papers
@router.post("/api/download-doc")
async def download_doc(request: DownloadDocRequest):
    try:
        if not request.papers or len(request.papers) == 0:
            raise HTTPException(status_code=400, detail="No papers to include in the document")

        # Create a new Document object
        doc = Document()

        # Add a title
        doc.add_heading(f"Summarized Papers for Topic: {request.topic}", level=1)

        # Add a table for the papers
        table = doc.add_table(rows=1, cols=4)
        table.autofit = True

        # Add table headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Paper Name'
        hdr_cells[1].text = 'Publication Year'
        hdr_cells[2].text = 'Author Names'
        hdr_cells[3].text = 'Summary'

        # Style the table headers
        for cell in hdr_cells:
            cell.font = doc.styles['Normal'].font
            cell.font.size = Pt(11)
            shading_elm = parse_xml(r'<w:shd {} w:fill="CCCCCC"/>'.format(nsdecls('w')))
            cell._element.get_or_add_tcPr().append(shading_elm)

        # Add each summarized paper to the table
        for paper in request.papers:
            row_cells = table.add_row().cells
            row_cells[0].text = paper.title
            row_cells[1].text = str(paper.year)
            row_cells[2].text = paper.authors
            row_cells[3].text = paper.summary

        # Save the document as a DOCX file
        file_path = f"summarized_papers_{request.topic}.docx"
        doc.save(file_path)

        # Return the file as a response for download
        return FileResponse(file_path, filename=file_path)

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

app.include_router(router, prefix="/lit", tags=["Literature Review Generator"])